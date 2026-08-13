"""Сохранение расшифровки в Word (.docx) и в обычный текст."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from engine import Line, Result, format_time


@dataclass
class Turn:
    """Слитая реплика: подряд идущие куски одного говорящего."""

    speaker: str | None
    start: float
    end: float
    text: str


def speaker_blocks(lines: list[Line]) -> list[tuple[str | None, list[Line]]]:
    """Собирает подряд идущие реплики одного человека в блок.

    В отличие от group_turns текст не склеивается: внутри блока каждая реплика
    остаётся отдельной строкой со своим временем — так устроена расшифровка в Word.
    """
    blocks: list[tuple[str | None, list[Line]]] = []
    for line in lines:
        if blocks and blocks[-1][0] == line.speaker:
            blocks[-1][1].append(line)
        else:
            blocks.append((line.speaker, [line]))
    return blocks


def stamp(seconds: float) -> str:
    """Время в виде 00:02:19 — с часами всегда, как принято в расшифровках."""
    seconds = max(int(seconds or 0), 0)
    hours, rest = divmod(seconds, 3600)
    minutes, secs = divmod(rest, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def group_turns(lines: list[Line]) -> list[Turn]:
    """Склеивает соседние куски одного говорящего в цельные реплики.

    Whisper режет речь по паузам, и без склейки документ превращается
    в лесенку из огрызков по три слова.
    """
    turns: list[Turn] = []
    for line in lines:
        if turns and turns[-1].speaker == line.speaker:
            turns[-1].text = (turns[-1].text + " " + line.text).strip()
            turns[-1].end = line.end
        else:
            turns.append(Turn(line.speaker, line.start, line.end, line.text))
    return turns


def render_text(results: list[Result], show_time: bool, show_speaker: bool) -> str:
    """Собирает расшифровку в обычный текст — то же, что видно в окне."""
    blocks: list[str] = []
    for result in results:
        pieces: list[str] = []
        if len(results) > 1:
            pieces.append(f"── {Path(result.path).name} ──\n")

        # Тот же вид, что и в документе Word: имя говорящего строкой,
        # под ним реплики со временем.
        for speaker, block in speaker_blocks(result.lines):
            if show_speaker and speaker:
                pieces.append(f"\n{speaker}:")
            for line in block:
                pieces.append(f"{stamp(line.start)} - {line.text}" if show_time else line.text)

        blocks.append("\n".join(pieces))
    return "\n\n".join(blocks).strip()


def save_txt(path: Path, results: list[Result], show_time: bool, show_speaker: bool) -> None:
    path.write_text(render_text(results, show_time, show_speaker) + "\n", encoding="utf-8")


def save_docx(path: Path, results: list[Result], show_time: bool, show_speaker: bool) -> None:
    """Складывает документ Word в виде обычной расшифровки.

    Имя говорящего — отдельной строкой, под ним реплики вида «00:02:19 - текст»,
    абзацы выключены по ширине.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.15

    for index, result in enumerate(results):
        if index:
            document.add_page_break()

        document.add_heading(Path(result.path).name, level=1)

        facts = []
        if result.lines:
            facts.append("длительность " + format_time(result.lines[-1].end))
        if result.language:
            facts.append("язык " + result.language)
        if result.speakers:
            facts.append(f"говорящих: {result.speakers}")
        if facts:
            note = document.add_paragraph(" · ".join(facts))
            note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        for speaker, block in speaker_blocks(result.lines):
            if show_speaker and speaker:
                title = document.add_paragraph()
                title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                title.paragraph_format.space_before = Pt(16)
                title.paragraph_format.space_after = Pt(4)
                title.add_run(f"{speaker}:").bold = True

            for line in block:
                paragraph = document.add_paragraph()
                if show_time:
                    paragraph.add_run(f"{stamp(line.start)} - ")
                paragraph.add_run(line.text)

    document.save(str(path))
